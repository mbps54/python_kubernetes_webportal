
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_ROW_HEIGHT_RULE

'''
data = {'email': 'i.ivanov@akkuyu.comn',
        'name1': 'Иванова',
        'name2': 'Ивана',
        'name3': 'Ивановича',
        'shop': 'ЦПЧ',
        'phone': '9922',
        'kimlik': '99123456789',
        'rabotnik_name1': 'Иванов',
        'rabotnik_name2': 'Иван',   
        'rabotnik_name3': 'Иванович',
        'rabotnik_name_lat1': 'Ivan',
        'rabotnik_name_lat2': 'Ivanov',
        'rabotnik_birthdate': '01.01.2000',
        'rabotnik_zagran': '99 44332211',
        'rabotnik_zagran_srok': '01.01.2030',
        'supruga_name1': 'Николаева',
        'supruga_name2': 'Ольга',
        'supruga_name3': 'Милайловна',
        'supruga_name_lat1': 'Olga',
        'supruga_name_lat2': 'Mikhailovna',
        'supruga_birthdate': '02.02.2002',
        'supruga_zagran': '99 11223344',
        'supruga_zagran_srok': '01.01.2030',
        'supruga_brak_number': '99873VX72',
        'supruga_brak_data': '01.01.2020',
        'child1_name1': 'Иванов',
        'child1_name2': 'Никита',
        'child1_name3': 'Иванович',
        'child1_name_lat1': 'Nikita',
        'child1_name_lat2': 'Ivanov',
        'child1_birthdate': '01.05.2020',
        'child1_zagran': '99 11223344',
        'child1_zagran_srok': '01.01.2024',
        'child1_svidetelstvo_number': '99873V-X72',
        'child1_svidetelstvo_data': '20.05.2020',
        'child2_name1': 'Иванова',
        'child2_name2': 'Марья',
        'child2_name3': 'Ивановна',
        'child2_name_lat1': 'Maria',
        'child2_name_lat2': 'Ivanova',
        'child2_birthdate': '01.05.2021',
        'child2_zagran': '99 22334455',
        'child2_zagran_srok': '01.01.2025',
        'child2_svidetelstvo_number': '992323V-Y72',
        'child2_svidetelstvo_data': '20.03.2021',
        'child3_name1': '',
        'child4_name1': '',
        'child5_name1': '',
        'ticket_date': '20.12.2021',
        'hotel_date': '21.12.2021',
        'departure': 'Санкт-Петербург',
        'arrival': 'Адана',
        'hotel': 'Park Inn Tasucu by Radisson',
        }
'''

def set_col_widths(table):
    widths = (Cm(2.3), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.4), Cm(2.7))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)


def set_col_widths_2(table):
    widths = (Cm(8.4), Cm(8.4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)

def set_col_widths_3(table):
    widths = (Cm(3.5), Cm(7.5), Cm(5.8))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
        row.height = Cm(1.4)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
#        row.height = Cm(1.3)
#        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 0
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)
                    font.name= 'Arial Narrow'


def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table

def doc1(data):


    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


    if data['supruga_name1'] != '' and data['child1_name1'] != '':
        wifechildren = ', и следующим членам моей семьи (супруг/супруга и несовершеннолетние дети)'
    elif data['supruga_name1'] != '' and data['child1_name1'] == '':
        wifechildren = ', и следующим членам моей семьи (супруг/супруга)'
    elif data['supruga_name1'] == '' and data['child1_name1'] != '':
        wifechildren = ', и следующим членам моей семьи (несовершеннолетние дети)'

    p1 = document.add_paragraph('Директору по персоналу')
    p1.alignment = 2
    p1.style = style
    p1.add_run('\nAKKUYU NÜKLEER ANONİM ŞİRKETİ')
    p1.add_run('\nА.А. Павлюку')
    p1.add_run('\nОт {} {} {}'.format(data['name1'], data['name2'], data['name3']))
    p1.add_run('\n{}'.format(data['shop']))
    p1.add_run('\nКонтактный телефон {}'.format(data['phone']))
    p1.add_run('\nE-mail {}'.format(data['email']))
    p1.add_run('\nID ВНЖ (кимлик) {}'.format(data['kimlik']))

    p2 = document.add_paragraph()
    p2.style = style
    p2.alignment = 1
    p2.add_run('ЗАЯВЛЕНИЕ').bold = True
    p2.paragraph_format.space_before = Pt(0)

    p3 = document.add_paragraph('')
    p3.style = style
    p3.add_run('                В соответствии с п.3.2.1 Политики по релокации персонала ').bold = False
    p3.add_run('AKKUYU NÜKLEER A. Ş.').bold = False
    p3.alignment = 0
    p3.add_run(
                f', прошу предоставить мне оплачиваемый билет на самолет/поезд'
                f' к месту моей работы{wifechildren}:'
              ).bold = False

    headers = ('Статус члена семьи', 'Ф.И.О.', 'First Name,\nLast Name',
               'Дата рождения', 'N загран-\nпаспорта', 'Срок действия\nзагранпаспорта')
    headers2 = ('ГОРОД ОТПРАВЛЕНИЯ', 'ГОРОД ПРИБЫТИЯ')
    records_table_2 = [[data['departure'], data['arrival']]]

    records_table = []
    if data['rabotnik_name1'] != '':
        status = 'Работник'
        fio = str(data['rabotnik_name1'] + '\n' + data['rabotnik_name2'] + '\n' + data['rabotnik_name3'])
        firstlastname = str(data['rabotnik_name_lat1'] +'\n' + data['rabotnik_name_lat2'])
        birthday = data['rabotnik_birthdate']
        zagran = data['rabotnik_zagran']
        zagran_srok = data['rabotnik_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['supruga_name1'] != '':
        status = 'Супруг/\nсупруга'
        fio = str(data['supruga_name1'] + '\n' + data['supruga_name2'] + '\n' + data['supruga_name3'])
        firstlastname = str(data['supruga_name_lat1'] +'\n' + data['supruga_name_lat2'])
        birthday = data['supruga_birthdate']
        zagran = data['supruga_zagran']
        zagran_srok = data['supruga_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child1_name1'] != '':
        status = 'Ребенок 1'
        fio = str(data['child1_name1'] + '\n' + data['child1_name2'] + '\n' + data['child1_name3'])
        firstlastname = str(data['child1_name_lat1'] +'\n' + data['child1_name_lat2'])
        birthday = data['child1_birthdate']
        zagran = data['child1_zagran']
        zagran_srok = data['child1_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child2_name1'] != '':
        status = 'Ребенок 2'
        fio = str(data['child2_name1'] + '\n' + data['child2_name2'] + '\n' + data['child2_name3'])
        firstlastname = str(data['child2_name_lat1'] +'\n' + data['child2_name_lat2'])
        birthday = data['child2_birthdate']
        zagran = data['child2_zagran']
        zagran_srok = data['child2_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child3_name1'] != '':
        status = 'Ребенок 3'
        fio = str(data['child3_name1'] + '\n' + data['child3_name2'] + '\n' + data['child3_name3'])
        firstlastname = str(data['child3_name_lat1'] +'\n' + data['child3_name_lat2'])
        birthday = data['child3_birthdate']
        zagran = data['child3_zagran']
        zagran_srok = data['child3_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child4_name1'] != '':
        status = 'Ребенок 4'
        fio = str(data['child4_name1'] + '\n' + data['child4_name2'] + '\n' + data['child4_name3'])
        firstlastname = str(data['child4_name_lat1'] +'\n' + data['child4_name_lat2'])
        birthday = data['child4_birthdate']
        zagran = data['child4_zagran']
        zagran_srok = data['child4_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child5_name1'] != '':
        status = 'Ребенок 5'
        fio = str(data['child5_name1'] + '\n' + data['child5_name2'] + '\n' + data['child5_name3'])
        firstlastname = str(data['child5_name_lat1'] +'\n' + data['child5_name_lat2'])
        birthday = data['child5_birthdate']
        zagran = data['child5_zagran']
        zagran_srok = data['child5_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)

    table1 = create_table(document, headers, records_table)

    table1.allow_autofit = False
    table1.autofit = False
    set_col_widths(table1)

    p4 = document.add_paragraph()
    p4.style = style
    ticketdate = data['ticket_date']
    p4.add_run(f'Авиабилеты прошу приобрести на ')
    p4.add_run(f'{ticketdate}').bold = True
    p4.paragraph_format.space_before = Pt(10)

    table2 = create_table(document, headers2, records_table_2)
    set_col_widths_2(table2)

    hotel = data['hotel']
    hoteldate = data['hotel_date']
    p5 = document.add_paragraph()
    p5.style = style
    p5.add_run('                В соответствии с п.3.2.3 Политики по релокации персонала ')
    p5.add_run('AKKUYU NÜKLEER ANONİM ŞİRKETİ ')
    p5.alignment = 0
    p5.add_run(f' прошу забронировать гостиницу ')
    p5.add_run(f'{hotel}').bold = True
    p5.add_run(f'.\n')
    p5.add_run(f'Дата заселения в гостиницу ')
    p5.add_run(f'{hoteldate}').bold = True
    p5.add_run(f' до момента переезда на постоянную квартиру, но не более 30 календарных дней в рамках '
               f'установленных лимитов.')
    p5.paragraph_format.space_before = Pt(10)

    if data['supruga_brak_number'] != '' or data['child1_svidetelstvo_data'] != '':
        p6 = document.add_paragraph()
        tmp = p6.add_run('Приложения (при необходимости, для подтверждения права предоставления льготы):')
        tmp.font.size = Pt(11)
        tmp.italic = True
        p6.alignment = 0
        if data['supruga_brak_number'] != '':
            suprugabraknumber = data['supruga_brak_number']
            suprugabrakdate = data['supruga_brak_data']
            tmp = p6.add_run(
                        f'\n    1. Свидетельство о регистрации брака № {suprugabraknumber}, дата {suprugabrakdate}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child1_name1'] != '':
            if data['supruga_brak_number'] != '':
                i = '    2'
            else:
                i = '    1'
            child1svidetelstvonumber = data['child1_svidetelstvo_number']
            child1svidetelstvodata = data['child1_svidetelstvo_data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child2_name1'] != '':
            if data['supruga_brak_number'] != '':
                i = '    3'
            else:
                i = '    2'
            child1svidetelstvonumber = data['child2_svidetelstvo_number']
            child1svidetelstvodata = data['child2_svidetelstvo_data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child3_name1'] != '':
            if data['supruga_brak_number'] != '':
                i = '    4'
            else:
                i = '    3'
            child1svidetelstvonumber = data['child3_svidetelstvo_number']
            child1svidetelstvodata = data['child3_svidetelstvo_data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child4_name1'] != '':
            if data['supruga_brak_number'] != '':
                i = '    5'
            else:
                i = '    4'
            child1svidetelstvonumber = data['child4_svidetelstvo_number']
            child1svidetelstvodata = data['child4_svidetelstvo_data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child5_name1'] != '':
            if data['supruga_brak_number'] != '':
                i = '    6'
            else:
                i = '    5'
            child1svidetelstvonumber = data['child2_svidetelstvo_number']
            child1svidetelstvodata = data['child2_svidetelstvo_data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        p6.paragraph_format.space_before = Pt(10)

        p7 = document.add_paragraph()
        p7.add_run('Подтверждаю достоверность документов, обосновывающих семейный статус.').italic = False
        p7.paragraph_format.space_before = Pt(10)

    fio = data['rabotnik_name1'] + ' ' + data['rabotnik_name2'] + ' ' + data['rabotnik_name3']
    p8 = document.add_paragraph()
    currdate = datetime.today().strftime('%d.%m.%Y')
    p8.add_run(f'                             / {fio} / {currdate}\n').underline = True
    tmp = p8.add_run(f'                 Подпись, расшифровака, дата                          ')
    tmp.font.size = Pt(9)
    tmp.italic = True
    p8.paragraph_format.space_before = Pt(13)
    p8.alignment = 2


    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(1.7)


    section = document.sections[0]
    footer = section.footer
    table3 = footer.add_table(0, 3, Cm(6))

    row_cells = table3.add_row().cells
    row_cells[0].text = str('\nPL-01-04-06/05-01/01')
    row_cells[1].text = str('Первый въезд иностранного работника (билеты, гостиница, виза, багаж, подъемные, переезд семьи, компенсация НДФЛ)')
    row_cells[2].text = str('Члены семьи работника: супруг/супруга, действующий законный брак, несовершеннолетние дети до 18 лет')
    table3.allow_autofit = False
    table3.autofit = False
    set_col_widths_3(table3)
    table3.allow_autofit = False
    table3.autofit = False
    table3.style = 'Table Grid'


    document.save('./files/doc1/document.docx')

#doc1(data)

from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_ROW_HEIGHT_RULE

data = {'email': 'i.ivanov@akkuyu.comn',
        'name1': 'Иванова',
        'name2': 'Ивана',
        'name3': 'Ивановича',
        'shop': 'ЦПЧ',
        'phone': '9922',
        'kimlik': '99123456789',

        'rabotnik-name1': 'Иванов',
        'rabotnik-name2': 'Иван',
        'rabotnik-name3': 'Иванович',
        'rabotnik-name_lat1': 'Ivan',
        'rabotnik-name_lat2': 'Ivanov',
        'rabotnik-birthdate': '01.01.2000',
        'rabotnik-zagran': '99 44332211',
        'rabotnik-zagran-srok': '01.01.2030',

        'supruga-name1': 'Николаева',
        'supruga-name2': 'Ольга',
        'supruga-name3': 'Милайловна',
        'supruga-name_lat1': 'Olga',
        'supruga-name_lat2': 'Mikhailovna',
        'supruga-birthdate': '02.02.2002',
        'supruga-zagran': '99 11223344',
        'supruga-zagran-srok': '01.01.2030',
        'supruga-brak-number': '99873VX72',
        'supruga-brak-data': '01.01.2020',

        'child1-name1': 'Иванов',
        'child1-name2': 'Никита',
        'child1-name3': 'Иванович',
        'child1-name_lat1': 'Nikita',
        'child1-name_lat2': 'Ivanov',
        'child1-birthdate': '01.05.2020',
        'child1-zagran': '99 11223344',
        'child1-zagran-srok': '01.01.2024',
        'child1-svidetelstvo-number': '99873V-X72',
        'child1-svidetelstvo-data': '20.05.2020',

        'child2-name1': 'Иванова',
        'child2-name2': 'Марья',
        'child2-name3': 'Ивановна',
        'child2-name_lat1': 'Maria',
        'child2-name_lat2': 'Ivanova',
        'child2-birthdate': '01.05.2021',
        'child2-zagran': '99 22334455',
        'child2-zagran-srok': '01.01.2025',
        'child2-svidetelstvo-number': '992323V-Y72',
        'child2-svidetelstvo-data': '20.03.2021',

        'child3-name1': '',
        'child4-name1': '',
        'child5-name1': '',

        'ticket-date': '20.12.2021',
        'hotel-date': '21.12.2021',
        'departure': 'Санкт-Петербург',
        'arrival': 'Адана',
        'hotel': 'Park Inn Tasucu by Radisson',

}

'''
def set_col_widths(table):
    widths = (Inches(1), Inches(2), Inches(1.5))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
'''

def set_col_widths(table):
    widths = (Cm(2.3), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.4), Cm(2.7))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(10)


def set_col_widths_2(table):
    widths = (Cm(8.4), Cm(8.4))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(12)

def set_col_widths_3(table):
    widths = (Cm(3.5), Cm(7.5), Cm(5.8))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width
        row.height = Cm(1.4)
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
#    for row in table.rows:
#        for idx, col in enumerate(table.columns):
#            col.width = widths[idx]
#        row.height = Cm(1.3)
#        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                paragraph.alignment = 0
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(9)
                    font.name= 'Arial Narrow'


def create_table(document, headers, rows, style='Table Grid'):
    cols_number = len(headers)

    table = document.add_table(rows=1, cols=cols_number)
    table.style = style

    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]

    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])

    return table

def doc1(data):
    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)


    if data['supruga-name1'] != '' and data['child1-name1'] != '':
        wifechildren = ', и следующим членам моей семьи (супруг/супруга и несовершеннолетние дети)'
    elif data['supruga-name1'] != '' and data['child1-name1'] == '':
        wifechildren = ', и следующим членам моей семьи (супруг/супруга)'
    elif data['supruga-name1'] == '' and data['child1-name1'] != '':
        wifechildren = ', и следующим членам моей семьи (несовершеннолетние дети)'

    p1 = document.add_paragraph('Директору по персоналу')
    p1.alignment = 2
    p1.style = style
    p1.add_run('\nAKKUYU NÜKLEER ANONİM ŞİRKETİ')
    p1.add_run('\nА.А. Павлюку')
    p1.add_run('\nОт {} {} {}'.format(data['name1'], data['name2'], data['name3']))
    p1.add_run('\n{}'.format(data['shop']))
    p1.add_run('\nКонтактный телефон {}'.format(data['phone']))
    p1.add_run('\nE-mail {}'.format(data['email']))
    p1.add_run('\nID ВНЖ (кимлик) {}'.format(data['kimlik']))

    p2 = document.add_paragraph()
    p2.style = style
    p2.alignment = 1
    p2.add_run('ЗАЯВЛЕНИЕ').bold = True
    p2.paragraph_format.space_before = Pt(0)

    p3 = document.add_paragraph('')
    p3.style = style
    p3.add_run('                В соответствии с п.3.2.1 Политики по релокации персонала ').bold = False
    p3.add_run('AKKUYU NÜKLEER A. Ş.').bold = False
    p3.alignment = 0
    p3.add_run(
                f', прошу предоставить мне оплачиваемый билет на самолет/поезд'
                f' к месту моей работы{wifechildren}:'
              ).bold = False

    headers = ('Статус члена семьи', 'Ф.И.О.', 'First Name,\nLast Name',
               'Дата рождения', 'N загран-\nпаспорта', 'Срок действия\nзагранпаспорта')
    headers2 = ('ГОРОД ОТПРАВЛЕНИЯ', 'ГОРОД ПРИБЫТИЯ')
    records_table_2 = [[data['departure'], data['arrival']]]

    records_table = []
    if data['rabotnik-name1'] != '':
        status = 'Работник'
        fio = str(data['rabotnik-name1'] + '\n' + data['rabotnik-name2'] + '\n' + data['rabotnik-name3'])
        firstlastname = str(data['rabotnik-name_lat1'] +'\n' + data['rabotnik-name_lat2'])
        birthday = data['rabotnik-birthdate']
        zagran = data['rabotnik-zagran']
        zagran_srok = data['rabotnik-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['supruga-name1'] != '':
        status = 'Супруг/\nсупруга'
        fio = str(data['supruga-name1'] + '\n' + data['supruga-name2'] + '\n' + data['supruga-name3'])
        firstlastname = str(data['supruga-name_lat1'] +'\n' + data['supruga-name_lat2'])
        birthday = data['supruga-birthdate']
        zagran = data['supruga-zagran']
        zagran_srok = data['supruga-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child1-name1'] != '':
        status = 'Ребенок 1'
        fio = str(data['child1-name1'] + '\n' + data['child1-name2'] + '\n' + data['child1-name3'])
        firstlastname = str(data['child1-name_lat1'] +'\n' + data['child1-name_lat2'])
        birthday = data['child1-birthdate']
        zagran = data['child1-zagran']
        zagran_srok = data['child1-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child2-name1'] != '':
        status = 'Ребенок 2'
        fio = str(data['child2-name1'] + '\n' + data['child2-name2'] + '\n' + data['child2-name3'])
        firstlastname = str(data['child2-name_lat1'] +'\n' + data['child2-name_lat2'])
        birthday = data['child2-birthdate']
        zagran = data['child2-zagran']
        zagran_srok = data['child2-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child3-name1'] != '':
        status = 'Ребенок 3'
        fio = str(data['child3-name1'] + '\n' + data['child3-name2'] + '\n' + data['child3-name3'])
        firstlastname = str(data['child3-name_lat1'] +'\n' + data['child3-name_lat2'])
        birthday = data['child3-birthdate']
        zagran = data['child3-zagran']
        zagran_srok = data['child3-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child4-name1'] != '':
        status = 'Ребенок 4'
        fio = str(data['child4-name1'] + '\n' + data['child4-name2'] + '\n' + data['child4-name3'])
        firstlastname = str(data['child4-name_lat1'] +'\n' + data['child4-name_lat2'])
        birthday = data['child4-birthdate']
        zagran = data['child4-zagran']
        zagran_srok = data['child4-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)
    if data['child5-name1'] != '':
        status = 'Ребенок 5'
        fio = str(data['child5-name1'] + '\n' + data['child5-name2'] + '\n' + data['child5-name3'])
        firstlastname = str(data['child5-name_lat1'] +'\n' + data['child5-name_lat2'])
        birthday = data['child5-birthdate']
        zagran = data['child5-zagran']
        zagran_srok = data['child5-zagran-srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table.append(i)

    table1 = create_table(document, headers, records_table)

    table1.allow_autofit = False
    table1.autofit = False
    set_col_widths(table1)

    p4 = document.add_paragraph()
    p4.style = style
    ticketdate = data['ticket-date']
    p4.add_run(f'Авиабилеты прошу приобрести на ')
    p4.add_run(f'{ticketdate}').bold = True
    p4.paragraph_format.space_before = Pt(10)

    table2 = create_table(document, headers2, records_table_2)
    set_col_widths_2(table2)

    hotel = data['hotel']
    hoteldate = data['hotel-date']
    p5 = document.add_paragraph()
    p5.style = style
    p5.add_run('                В соответствии с п.3.2.3 Политики по релокации персонала ')
    p5.add_run('AKKUYU NÜKLEER ANONİM ŞİRKETİ ')
    p5.alignment = 0
    p5.add_run(f' прошу забронировать гостиницу ')
    p5.add_run(f'{hotel}').bold = True
    p5.add_run(f'.\n')
    p5.add_run(f'Дата заселения в гостиницу ')
    p5.add_run(f'{hoteldate}').bold = True
    p5.add_run(f' до момента переезда на постоянную квартиру, но не более 30 календарных дней в рамках '
               f'установленных лимитов.')
    p5.paragraph_format.space_before = Pt(10)

    if data['supruga-brak-number'] != '' or data['child1-svidetelstvo-data'] != '':
        p6 = document.add_paragraph()
        tmp = p6.add_run('Приложения (при необходимости, для подтверждения права предоставления льготы):')
        tmp.font.size = Pt(11)
        tmp.italic = True
        p6.alignment = 0
        if data['supruga-brak-number'] != '':
            suprugabraknumber = data['supruga-brak-number']
            suprugabrakdate = data['supruga-brak-data']
            tmp = p6.add_run(
                        f'\n    1. Свидетельство о регистрации брака № {suprugabraknumber}, дата {suprugabrakdate}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child1-name1'] != '':
            if data['supruga-brak-number'] != '':
                i = '    2'
            else:
                i = '    1'
            child1svidetelstvonumber = data['child1-svidetelstvo-number']
            child1svidetelstvodata = data['child1-svidetelstvo-data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child2-name1'] != '':
            if data['supruga-brak-number'] != '':
                i = '    3'
            else:
                i = '    2'
            child1svidetelstvonumber = data['child2-svidetelstvo-number']
            child1svidetelstvodata = data['child2-svidetelstvo-data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child3-name1'] != '':
            if data['supruga-brak-number'] != '':
                i = '    4'
            else:
                i = '    3'
            child1svidetelstvonumber = data['child3-svidetelstvo-number']
            child1svidetelstvodata = data['child3-svidetelstvo-data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child4-name1'] != '':
            if data['supruga-brak-number'] != '':
                i = '    5'
            else:
                i = '    4'
            child1svidetelstvonumber = data['child4-svidetelstvo-number']
            child1svidetelstvodata = data['child4-svidetelstvo-data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        if data['child5-name1'] != '':
            if data['supruga-brak-number'] != '':
                i = '    6'
            else:
                i = '    5'
            child1svidetelstvonumber = data['child2-svidetelstvo-number']
            child1svidetelstvodata = data['child2-svidetelstvo-data']
            tmp = p6.add_run(
                        f'\n{i}. Свидетельство о рождении ребенка № '
                        f'{child1svidetelstvonumber}, дата {child1svidetelstvodata}'
                      )
            tmp.font.size = Pt(11)
            tmp.italic = True

        p6.paragraph_format.space_before = Pt(10)

        p7 = document.add_paragraph()
        p7.add_run('Подтверждаю достоверность документов, обосновывающих семейный статус.').italic = False
        p7.paragraph_format.space_before = Pt(10)

    fio = data['rabotnik-name1'] + ' ' + data['rabotnik-name2'] + ' ' + data['rabotnik-name3']
    p8 = document.add_paragraph()
    currdate = datetime.today().strftime('%d.%m.%Y')
    p8.add_run(f'                             / {fio} / {currdate}\n').underline = True
    tmp = p8.add_run(f'                 Подпись, расшифровака, дата                          ')
    tmp.font.size = Pt(9)
    tmp.italic = True
    p8.paragraph_format.space_before = Pt(13)
    p8.alignment = 2


    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(1.7)


    section = document.sections[0]
    footer = section.footer
    table3 = footer.add_table(0, 3, Cm(6))

    row_cells = table3.add_row().cells
    row_cells[0].text = str('\nPL-01-04-06/05-01/01')
    row_cells[1].text = str('Первый въезд иностранного работника (билеты, гостиница, виза, багаж, подъемные, переезд семьи, компенсация НДФЛ)')
    row_cells[2].text = str('Члены семьи работника: супруг/супруга, действующий законный брак, несовершеннолетние дети до 18 лет')
    table3.allow_autofit = False
    table3.autofit = False
    set_col_widths_3(table3)
    table3.allow_autofit = False
    table3.autofit = False
    table3.style = 'Table Grid'


    document.save('./files/doc1/document.docx')

doc1(data)
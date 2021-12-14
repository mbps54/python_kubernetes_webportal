#!/usr/bin/python3
########################### IMPORT GENERAL MODULES  ############################
import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL

###########################      IMPORT MODULES     ############################
from documents_functions import create_table
from documents_functions import set_table_params_lo

###########################   CREATE DOC FUNCTION   ############################
def create_doc_1_pdf(data: dict) -> None:
    document = Document()

###########################      DEFAULT STYLE      ############################
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

###########################        TEXT IFs         ############################
    text = ''
    if data['rabotnik_zagran'] != '':
        if data['supruga_name1'] != '' and data['child1_name1'] != '':
            text = ', и следующим членам моей семьи (супруг/супруга и несовершеннолетние дети)'
        elif data['supruga_name1'] != '' and data['child1_name1'] == '':
            text = ', и следующим членам моей семьи (супруг/супруга)'
        elif data['supruga_name1'] == '' and data['child1_name1'] != '':
            text = ', и следующим членам моей семьи (несовершеннолетние дети)'
    else:
        if data['supruga_name1'] != '' and data['child1_name1'] != '':
            text = ' следующим членам моей семьи (супруг/супруга и несовершеннолетние дети)'
        elif data['supruga_name1'] != '' and data['child1_name1'] == '':
            text = ' следующим членам моей семьи (супруг/супруга)'
        elif data['supruga_name1'] == '' and data['child1_name1'] != '':
            text = ' следующим членам моей семьи (несовершеннолетние дети)'

###############################    PARAGRAPH 1   ###############################
    p1 = document.add_paragraph()
    p1.style = style
    p1.alignment = 2
    p1.add_run(f"Директору по персоналу"
               f"\nAKKUYU NÜKLEER ANONİM ŞİRKETİ"
               f"\nА.А. Павлюку"
               f"\nОт {(data['name1'])} {data['name2']} {data['name3']}"
               f"\n{data['shop']}"
               f"\nКонтактный телефон {data['phone']}"
               f"\nE-mail {data['email']}"
               f"\nID ВНЖ (kimlik) {data['kimlik']}"
    )

###############################    PARAGRAPH 2   ###############################
    p2 = document.add_paragraph()
    p2.style = style
    p2.alignment = 1
    p2.paragraph_format.space_before = Pt(0)
    p2.add_run('ЗАЯВЛЕНИЕ').bold = True

###############################    PARAGRAPH 3   ###############################
    p3 = document.add_paragraph()
    p3.alignment = 0
    p3.add_run(f"                "
               f"В соответствии с п.3.2.1 Политики по релокации персонала "
               f"AKKUYU NÜKLEER A. Ş."
    )
    if data['rabotnik_zagran'] != '':
        p3.add_run(f", прошу предоставить мне оплачиваемый билет на самолет/поезд"
                   f" к месту моей работы{text}:"
        )
    else:
        p3.add_run(f", прошу предоставить оплачиваемый билет на самолет/поезд"
                   f" к месту моей работы{text}:"
        )

###############################      TABLE1      ###############################
    headers_1 = (
        'Статус члена семьи',
        'Ф.И.О.',
        'First Name,\nLast Name',
        'Дата рождения',
        'N загран-\nпаспорта',
        'Срок действия\nзагранпаспорта',
    )
    records_table_1 = []
    if data['rabotnik_zagran'] != '':
        status = 'Работник'
        fio = str(
            data['rabotnik_name1']
            + '\n'
            + data['rabotnik_name2']
            + '\n'
            + data['rabotnik_name3']
        )
        firstlastname = str(
            data['rabotnik_name_lat1'] + '\n' + data['rabotnik_name_lat2']
        )
        birthday = data['rabotnik_birthdate']
        zagran = data['rabotnik_zagran']
        zagran_srok = data['rabotnik_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['supruga_name1'] != '':
        status = 'Супруг/\nсупруга'
        fio = str(
            data['supruga_name1']
            + '\n'
            + data['supruga_name2']
            + '\n'
            + data['supruga_name3']
        )
        firstlastname = str(
            data['supruga_name_lat1'] + '\n' + data['supruga_name_lat2']
        )
        birthday = data['supruga_birthdate']
        zagran = data['supruga_zagran']
        zagran_srok = data['supruga_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['child1_name1'] != '':
        status = 'Ребенок 1'
        fio = str(
            data['child1_name1']
            + '\n'
            + data['child1_name2']
            + '\n'
            + data['child1_name3']
        )
        firstlastname = str(data['child1_name_lat1'] + '\n' + data['child1_name_lat2'])
        birthday = data['child1_birthdate']
        zagran = data['child1_zagran']
        zagran_srok = data['child1_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['child2_name1'] != '':
        status = 'Ребенок 2'
        fio = str(
            data['child2_name1']
            + '\n'
            + data['child2_name2']
            + '\n'
            + data['child2_name3']
        )
        firstlastname = str(data['child2_name_lat1'] + '\n' + data['child2_name_lat2'])
        birthday = data['child2_birthdate']
        zagran = data['child2_zagran']
        zagran_srok = data['child2_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['child3_name1'] != '':
        status = 'Ребенок 3'
        fio = str(
            data['child3_name1']
            + '\n'
            + data['child3_name2']
            + '\n'
            + data['child3_name3']
        )
        firstlastname = str(data['child3_name_lat1'] + '\n' + data['child3_name_lat2'])
        birthday = data['child3_birthdate']
        zagran = data['child3_zagran']
        zagran_srok = data['child3_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['child4_name1'] != '':
        status = 'Ребенок 4'
        fio = str(
            data['child4_name1']
            + '\n'
            + data['child4_name2']
            + '\n'
            + data['child4_name3']
        )
        firstlastname = str(data['child4_name_lat1'] + '\n' + data['child4_name_lat2'])
        birthday = data['child4_birthdate']
        zagran = data['child4_zagran']
        zagran_srok = data['child4_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)
    if data['child5_name1'] != '':
        status = 'Ребенок 5'
        fio = str(
            data['child5_name1']
            + '\n'
            + data['child5_name2']
            + '\n'
            + data['child5_name3']
        )
        firstlastname = str(data['child5_name_lat1'] + '\n' + data['child5_name_lat2'])
        birthday = data['child5_birthdate']
        zagran = data['child5_zagran']
        zagran_srok = data['child5_zagran_srok']
        i = [status, fio, firstlastname, birthday, zagran, zagran_srok]
        records_table_1.append(i)

    table_1 = create_table(document, headers_1, records_table_1)
    table_1.allow_autofit = False
    table_1.autofit = False
    set_table_params_lo(table = table_1,
                     widths = (Cm(2.3), Cm(3.5), Cm(3.5), Cm(2.5), Cm(2.4), Cm(2.7)),
                     fontname = 'Times New Roman',
                     fontsize = Pt(10))

###############################    PARAGRAPH 4   ###############################
    p4 = document.add_paragraph()
    p4.style = style
    p4.paragraph_format.space_before = Pt(10)
    p4.add_run(f"Авиабилеты прошу приобрести на ")
    p4.add_run(f"{data['ticket_date']}").bold = True

###############################      TABLE2      ###############################
    headers_2 = ('ГОРОД ОТПРАВЛЕНИЯ', 'ГОРОД ПРИБЫТИЯ')
    records_table_2 = [[data['departure'], data['arrival']]]
    table_2 = create_table(document, headers_2, records_table_2)
    set_table_params_lo(table = table_2,
                     widths = (Cm(8.4), Cm(8.4)),
                     fontname = 'Times New Roman',
                     fontsize = Pt(12),
                     alignment = 1)

###############################    PARAGRAPH 5   ###############################
    if data['hotel_date'] != '' and data['hotel'] != '':
        p5 = document.add_paragraph()
        p5.style = style
        p5.alignment = 0
        p5.paragraph_format.space_before = Pt(10)
        p5.add_run('                ')
        p5.add_run('В соответствии с п.3.2.3 Политики по релокации персонала ')
        p5.add_run('AKKUYU NÜKLEER ANONİM ŞİRKETİ ')
        p5.add_run('прошу забронировать гостиницу ')
        p5.add_run(f"{data['hotel']}").bold = True
        p5.add_run('.\n')
        p5.add_run('Дата заселения в гостиницу ')
        p5.add_run(f"{data['hotel_date']} ").bold = True
        p5.add_run('до момента переезда на постоянную квартиру, но не более ')
        p5.add_run('30 календарных дней в рамках установленных лимитов.')

###############################    PARAGRAPH 6   ###############################
    if data['supruga_brak_number'] != '' or data['child1_svidetelstvo_data'] != '':
        p6 = document.add_paragraph()
        p6.alignment = 0
        p6.paragraph_format.space_before = Pt(10)
        tmp = p6.add_run('Приложения (для подтверждения права предоставления льготы):')
        tmp.font.size = Pt(11)
        tmp.italic = True
        if data['supruga_brak_number'] != '':
            tmp = p6.add_run(f"\n    1. Свидетельство о регистрации брака № "
                             f"{data['supruga_brak_number']}, "
                             f"дата {data['supruga_brak_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child1_name1'] != '':
            if data['supruga_brak_number'] != '':
                number = '    2'
            else:
                number = '    1'
            tmp = p6.add_run(
                f"\n{number}. Свидетельство о рождении ребенка № "
                f"{data['child1_svidetelstvo_number']}, "
                f"дата {data['child1_svidetelstvo_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child2_name1'] != '':
            if data['supruga_brak_number'] != '':
                number = '    3'
            else:
                number = '    2'
            tmp = p6.add_run(
                f"\n{number}. Свидетельство о рождении ребенка № "
                f"{data['child2_svidetelstvo_number']}, "
                f"дата {data['child2_svidetelstvo_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child3_name1'] != '':
            if data['supruga_brak_number'] != '':
                number = '    4'
            else:
                number = '    3'
            tmp = p6.add_run(
                f"\n{number}. Свидетельство о рождении ребенка № "
                f"{data['child3_svidetelstvo_number']}, "
                f"дата {data['child3_svidetelstvo_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child4_name1'] != '':
            if data['supruga_brak_number'] != '':
                number = '    5'
            else:
                number = '    4'
            tmp = p6.add_run(
                f"\n{number}. Свидетельство о рождении ребенка № "
                f"{data['child4_svidetelstvo_number']}, "
                f"дата {data['child4_svidetelstvo_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True
        if data['child5_name1'] != '':
            if data['supruga_brak_number'] != '':
                number = '    6'
            else:
                number = '    5'
            tmp = p6.add_run(
                f"\n{number}. Свидетельство о рождении ребенка № "
                f"{data['child5_svidetelstvo_number']}, "
                f"дата {data['child5_svidetelstvo_data']}"
            )
            tmp.font.size = Pt(11)
            tmp.italic = True

###############################    PARAGRAPH 7   ###############################
        p7 = document.add_paragraph()
        p7.paragraph_format.space_before = Pt(10)
        p7.add_run('Подтверждаю достоверность документов, обосновывающих семейный статус.')

###############################    PARAGRAPH 9   ###############################
    fio = (
        data['rabotnik_name1']
        + ' '
        + data['rabotnik_name2']
        + ' '
        + data['rabotnik_name3']
    )
    p8 = document.add_paragraph()
    p8.alignment = 2
    p8.paragraph_format.space_before = Pt(13)
    currdate = datetime.today().strftime('%d.%m.%Y')
    p8.add_run(f"                             / {fio} / {currdate}\n").underline = True
    tmp = p8.add_run(f"                 Подпись, расшифровака, дата          ")
    tmp.font.size = Pt(9)
    tmp.italic = True

###############################   PAGE BORDERS   ###############################
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(1.7)

###############################      FOOTER      ###############################
    section = document.sections[0]
    footer = section.footer
    table_3 = footer.add_table(0, 3, Cm(6))
    row_cells = table_3.add_row().cells
    row_cells[0].text = str('PL-01-04-06/05-01/01')
    row_cells[1].text = str('Первый въезд иностранного работника '
                            '(билеты, гостиница, виза, багаж, подъемные, '
                            'переезд семьи, компенсация НДФЛ)'
    )
    row_cells[2].text = str('Члены семьи работника: супруг/супруга, '
                            'действующий законный брак, несовершеннолетние '
                            'дети до 18 лет'
    )
    set_table_params_lo(table = table_3,
                     widths = (Cm(3.5), Cm(7.5), Cm(5.8)),
                     height = Cm(1.4),
                     heightrule = WD_ROW_HEIGHT_RULE.EXACTLY,
                     fontname = 'Nimbus Sans Narrow',
                     fontsize = Pt(9),
                     alignment = 0,
                     verticalalignment = WD_ALIGN_VERTICAL.CENTER)
    table_3.allow_autofit = False
    table_3.autofit = False
    table_3.style = 'Table Grid'

###############################    SAVE FILE     ###############################
    try:
        os.makedirs(os.path.expanduser('./files/doc1/'))
    except:
        pass
    document.save('./files/doc1/document_pdf.docx')

###############################  CONVERT TO PDF  ###############################
    subprocess.run(
        ["doc2pdf", "./files/doc1/document_pdf.docx"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
###############################    RENAME FILE   ###############################
    subprocess.run(
        ["mv", "./files/doc1/document_pdf.pdf",
         "./files/doc1/document.pdf"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
###############################        END       ###############################
    return(None)


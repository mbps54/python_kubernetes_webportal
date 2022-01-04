#!/usr/bin/python3
########################### IMPORT GENERAL MODULES  ############################
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL

###########################  CREATE TABLE FUNCTION  ############################
def create_table(document, headers, rows, style="Table Grid"):
    cols_number = len(headers)
    table = document.add_table(rows=1, cols=cols_number)
    table.style = style
    hdr_cells = table.rows[0].cells
    for i in range(cols_number):
        hdr_cells[i].text = headers[i]
    for row in rows:
        row_cells = table.add_row().cells
        for i in range(cols_number):
            row_cells[i].text = str(row[i])
    return table


###########################  TABLE PARAM FUNCTION   ############################
def set_table_params(
    table,
    widths=False,
    height=False,
    heightrule=False,
    fontname="Times New Roman",
    fontsize=Pt(12),
    alignment=0,
    verticalalignment=False,
):
    for row in table.rows:
        if widths:
            for idx, width in enumerate(widths):
                row.cells[idx].width = width
        if height:
            row.height = height
        if heightrule:
            row.height_rule = heightrule
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            if verticalalignment:
                cell.vertical_alignment = verticalalignment
            for paragraph in paragraphs:
                if alignment:
                    paragraph.alignment = alignment
                for run in paragraph.runs:
                    font = run.font
                    if fontname:
                        font.name = fontname
                    if fontsize:
                        font.size = fontsize


########################### TABLE PARAM FUNCTION LO ############################
def set_table_params_lo(
    table,
    widths=False,
    height=False,
    heightrule=False,
    fontname="Times New Roman",
    fontsize=Pt(12),
    alignment=0,
    verticalalignment=False,
):
    for row in table.rows:
        if widths:
            for idx, col in enumerate(table.columns):
                col.width = widths[idx]
        if height:
            row.height = height
        if heightrule:
            row.height_rule = heightrule
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            if verticalalignment:
                cell.vertical_alignment = verticalalignment
            for paragraph in paragraphs:
                if alignment:
                    paragraph.alignment = alignment
                for run in paragraph.runs:
                    font = run.font
                    if fontname:
                        font.name = fontname
                    if fontsize:
                        font.size = fontsize

###########################     ROUND FUNCTION      ############################
def round_a(number, ndigits=0):
    ndigits += 1
    n = round(number, ndigits)*(10**ndigits)
    m = n % 10
    n -= m
    if m >= 5:
        n += 10
    n /= (10**ndigits)
    if n % 1 == 0:
        return int(n)
    else:
        return n
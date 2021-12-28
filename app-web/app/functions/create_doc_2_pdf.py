#!/usr/bin/python3
########################### IMPORT GENERAL MODULES  ############################
import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL

###########################      IMPORT MODULES     ############################
from functions.documents_functions import create_table
from functions.documents_functions import set_table_params_lo

###########################   CREATE DOC FUNCTION   ############################
def create_doc_2_pdf(data: dict, result: dict) -> None:
    document = Document()

    ###########################      DEFAULT STYLE      ############################
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    ###############################    PARAGRAPH 1   ###############################
    p1 = document.add_paragraph()
    p1.style = style
    p1.alignment = 1
    p1.add_run("СТРУКТУРА ЗАРАБОТНОЙ ПЛАТЫ").bold = True

    ###############################    PARAGRAPH 2   ###############################
    p2 = document.add_paragraph()
    p2.style = style
    p2.alignment = 1
    p2.add_run(f"Исходные данные").bold = True

    ###############################      TABLE1      ###############################
    headers_1 = (
        'Параметр',
        'Значение',
    )
    records_table_1 = []
    records_table_1.append(['Оклад (в TRY)', f"{data['oklad']} TRY"])
    records_table_1.append(['ИСН (инд. стимулирующая надбавка)', f"{data['isn']} TRY"])
    records_table_1.append(['Доплата (совмещение и пр.)', f"{data['extra']} TRY"])
    records_table_1.append(['Целевой размер премии по КПЭ', f"{data['targetkpi']} TRY"])
    records_table_1.append(['Корректирующий коэффициент (Кк)', data["Kk"]])
    records_table_1.append(['Рассчетный коэффициент (Рк), влияние курса RUB/USD и Кк', result["Rk"]])
    records_table_1.append(['Валютный коэффициент (Кв), влияние курса TRY/USD и Рк', result["Kv"]])


    records_table_1.append(['Курс USD/RUB', data["CURRENT_USDRUB"]])
    records_table_1.append(['Курс USD/TRY', data["CURRENT_USDTRY"]])
    records_table_1.append(['Курс TRY/RUB', f"{data['CURRENT_TRYRUB']} (кросс-курс рассчитан автоматически)"])

    table_1 = create_table(document, headers_1, records_table_1)
    table_1.allow_autofit = False
    table_1.autofit = False
    set_table_params_lo(
        table=table_1,
        widths=(Cm(12), Cm(4)),
        fontname="Times New Roman",
        fontsize=Pt(12),
    )

    ###############################    PARAGRAPH 3   ###############################
    p3 = document.add_paragraph()
    p3.style = style
    p3.alignment = 1
    p3.add_run(f"Зарплата").bold = True
    p3.paragraph_format.space_before = Pt(12)

    ###############################      TABLE2      ###############################
    headers_2 = (
        'Раздел дохода',
        'У.Е.',
        'TRY',
        'RUB',
        'USD',
    )
    records_table_2 = []
    records_table_2.append(['Эквивалент заработной платы', result['ezp'], '', '', ''])
    records_table_2.append(['Оклад', '', data['oklad'], '', ''])
    records_table_2.append(['ИСН (инд. стимулирующая надбавка)', '', data['isn'], '', ''])
    records_table_2.append(['Индексирующая выплата', '', result["indincome"], '', ''])
    records_table_2.append(['Дополнительная доплата до эквивалента', '', result["zp_extra"], '', ''])
    records_table_2.append(['Итого к начислению', '', result["zp_TRY"], result["zp_RUB"], result["zp_USD"]])

    table_2 = create_table(document, headers_2, records_table_2)
    table_2.allow_autofit = False
    table_2.autofit = False
    set_table_params_lo(
        table=table_2,
        widths=(Cm(8), Cm(2), Cm(2), Cm(2), Cm(2)),
        fontname="Times New Roman",
        fontsize=Pt(12),
    )

    ###############################    PARAGRAPH 4   ###############################
    p4 = document.add_paragraph()
    p4.style = style
    p4.alignment = 1
    p4.add_run(f"Дополнительный доход (совмещение и пр.)").bold = True
    p4.paragraph_format.space_before = Pt(12)

    ###############################      TABLE 3      ###############################
    headers_3 = (
        'Раздел дохода',
        'У.Е.',
        'TRY',
        'RUB',
        'USD',
    )
    records_table_3 = []
    records_table_3.append(['Доплата (совмещение и пр.)', '', data['extra'], '', ''])
    records_table_3.append(['Итого к начислению', '', result["extra_2_try"], result["extra_2_rub"], result["extra_2_usd"]])

    table_3 = create_table(document, headers_3, records_table_3)
    table_3.allow_autofit = False
    table_3.autofit = False
    set_table_params_lo(
        table=table_3,
        widths=(Cm(8), Cm(2), Cm(2), Cm(2), Cm(2)),
        fontname="Times New Roman",
        fontsize=Pt(12),
    )

    ###############################    PARAGRAPH 5   ###############################
    p5 = document.add_paragraph()
    p5.style = style
    p5.alignment = 1
    p5.add_run(f"Премия").bold = True
    p5.paragraph_format.space_before = Pt(12)

    ###############################      TABLE 4      ###############################
    headers_4 = (
        'Раздел дохода',
        'У.Е.',
        'TRY',
        'RUB',
        'USD',
    )
    records_table_4 = []
    records_table_4.append(['Эквивалент целевого размера премии по КПЭ', result['ebonus'], '', '', ''])
    records_table_4.append(['Целевой размер премии по КПЭ', '', data['targetkpi'], '', ''])
    records_table_4.append(['Индексирующая выплата', '', result["indbonus"], '', ''])
    records_table_4.append(['Доплата до эквивалента', '', result["bonus_dop"], '', ''])
    records_table_4.append(['Итого к начислению', '', result["bonus_TRY"], result["bonus_RUB"], result["bonus_USD"]])

    table_4 = create_table(document, headers_4, records_table_4)
    table_4.allow_autofit = False
    table_4.autofit = False
    set_table_params_lo(
        table=table_4,
        widths=(Cm(8), Cm(2), Cm(2), Cm(2), Cm(2)),
        fontname="Times New Roman",
        fontsize=Pt(12),
    )


    ###############################   PAGE BORDERS   ###############################
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(1.7)



    ###############################    SAVE FILE     ###############################
    try:
        os.makedirs(os.path.expanduser("./files/doc2/"))
    except:
        pass
    document.save("./files/doc2/document_pdf.docx")

    ###############################  CONVERT TO PDF  ###############################
    subprocess.run(
        ["doc2pdf", "./files/doc2/document_pdf.docx"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
    ###############################    RENAME FILE   ###############################
    subprocess.run(
        ["mv", "./files/doc2/document_pdf.pdf", "./files/doc2/document.pdf"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
    ###############################        END       ###############################
    return None

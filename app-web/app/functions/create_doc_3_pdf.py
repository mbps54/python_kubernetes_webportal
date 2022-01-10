#!/usr/bin/python3
########################### IMPORT GENERAL MODULES  ############################
import os
import subprocess
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from datetime import datetime
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.table import WD_ALIGN_VERTICAL

###########################      IMPORT MODULES     ############################
from functions.documents_functions import create_table
from functions.documents_functions import set_table_params_lo

###########################   CREATE DOC FUNCTION   ############################
def create_doc_3_pdf(data: dict, user = '') -> None:
    document = Document()

###########################      DEFAULT STYLE      ############################
    style = document.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

###########################        TEXT IFs         ############################


###############################    PARAGRAPH 1   ###############################
    p1 = document.add_paragraph()
    p1.style = style
    p1.alignment = 2
    p1.add_run(
        f"Директору по персоналу"
        f"\nAKKUYU NÜKLEER ANONİM ŞİRKETİ"
        f"\nА.А. Павлюку"
        f"\nОт {(data['name1'])} {data['name2']} {data['name3']}"
        f"\n{data['shop']}"
        f"\nКонтактный телефон {data['phone']}"
        f"\nE-mail {data['email']}"
        f"\nID ВНЖ (kimlik) {data['kimlik']}"
    )

###############################    PARAGRAPH 2   ###############################
    p2 = document.add_paragraph()
    p2.style = style
    p2.alignment = 1
    p2.paragraph_format.space_before = Pt(0)
    p2.add_run("ЗАЯВЛЕНИЕ\n").bold = True
    p2.add_run("на выплату единовременного пособия и компенсацию визы, трансфера, багажа").bold = True

###############################    PARAGRAPH 3   ###############################
    p3 = document.add_paragraph()
    p3.alignment = 0
    p3.add_run(
        f"                "
        f"В соответствии с п.3.3.3 Политики по релокации персонала "
        f"AKKUYU NÜKLEER ANONİM ŞİRKETİ "
        f"прошу выплатить мне единовременное пособие в связи с началом работы в "
        f"AKKUYU NÜKLEER ANONİM ŞİRKETİ"
    )
    if data["child1_name1"] != "":
        data["approve_docs"] = True
        p3.add_run(
            f" и в связи с переездом моих детей к месту моей работы:\n"
        )
        p3.add_run(
            f"."
        )
    else:
        p3.add_run(
            f"."
        )
    if data["child1_name1"] != "":
        fio = str(
            data["child1_name1"]
            + " "
            + data["child1_name2"]
            + " "
            + data["child1_name3"]
        )
        p3.add_run(
            f"                "
            f"1) {fio}\n"
        )
    if data["child2_name1"] != "":
        fio = str(
            data["child2_name1"]
            + " "
            + data["child2_name2"]
            + " "
            + data["child2_name3"]
        )
        p3.add_run(
            f"                "
            f"2) {fio}\n"
        )
    if data["child3_name1"] != "":
        fio = str(
            data["child3_name1"]
            + " "
            + data["child3_name2"]
            + " "
            + data["child3_name3"]
        )
        p3.add_run(
            f"                "
            f"3) {fio}\n"
        )
    if data["child4_name1"] != "":
        fio = str(
            data["child4_name1"]
            + " "
            + data["child4_name2"]
            + " "
            + data["child4_name3"]
        )
        p3.add_run(
            f"                "
            f"4) {fio}\n"
        )
    if data["child5_name1"] != "":
        fio = str(
            data["child5_name1"]
            + " "
            + data["child5_name2"]
            + " "
            + data["child5_name3"]
        )
        p3.add_run(
            f"                "
            f"5) {fio}\n"
        )
        
###############################    PARAGRAPH 4   ###############################
    if data["lagguage_cost"] != "" or if data["transfer_cost"] != "":
        data["approve_docs"] = True
        p4 = document.add_paragraph()
        p4.alignment = 0
        p4.add_run(
            f"                "
            f"В соответствии с пп.3.2.1 и 3.2.2 Политики по релокации персонала "
            f"AKKUYU NÜKLEER ANONİM ŞİRKETİ прошу компенсировать мне расходы: "
        )
        if data["lagguage_cost"] != "":
            p4.add_run(
                f"                "
                f"1) по перевозке багажа в размере {data["lagguage_cost"]} "
                f"(указана итоговая сумма согласно подтверждающим документам)\n"
            )
        if data["transfer_cost"] != "":
            p4.add_run(
                f"                "
                f"2) на трансфер от {data["franster_from"]} (пункт отправления) "
                f"до {data["franster_to"]}(пункт назначения) "
                f"в размере {data["transfer_cost"]} "
                f"(указана итоговая сумма согласно подтверждающим документам)\n"
            )

###############################    PARAGRAPH 5   ###############################
    if data["visa_cost"] != "":
        data["approve_docs"] = True
        p5 = document.add_paragraph()
        p5.alignment = 0
        p5.add_run(
            f"                "
            f"В соответствии с п.3.1.2 Политики по релокации персонала "
            f"AKKUYU NÜKLEER ANONİM ŞİRKETİ прошу компенсировать мне расходы, "
            f"связанные с оформлением миграционных документов (визы) для моего "
            f"переезда к месту работы, в размере {data["visa_cost"]}"
            f"(указана итоговая сумма согласно подтверждающим документам)."
        )

###############################    PARAGRAPH 6   ###############################
    if data["approve_docs"] == True:
        p6 = document.add_paragraph()
        p6.alignment = 0
        p6.add_run(
            f"                "
            f"Подтверждающие документы прилагаю."
        )
    if data["child1_name1"] != "":
        fio = str(
            data["child1_name1"]
            + " "
            + data["child1_name2"]
            + " "
            + data["child1_name3"]
        )
        p6.add_run(
            f"                "
            f"- Свидетельство о рождении {fio}\n"
        )
    if data["child2_name1"] != "":
        fio = str(
            data["child2_name1"]
            + " "
            + data["child2_name2"]
            + " "
            + data["child2_name3"]
        )
        p6.add_run(
            f"                "
            f"- Свидетельство о рождении {fio}\n"
        )
    if data["child3_name1"] != "":
        fio = str(
            data["child3_name1"]
            + " "
            + data["child3_name2"]
            + " "
            + data["child3_name3"]
        )
        p3.add_run(
            f"                "
            f"- Свидетельство о рождении {fio}\n"
        )
    if data["child4_name1"] != "":
        fio = str(
            data["child4_name1"]
            + " "
            + data["child4_name2"]
            + " "
            + data["child4_name3"]
        )
        p3.add_run(
            f"                "
            f"- Свидетельство о рождении {fio}\n"
        )
    if data["child5_name1"] != "":
        fio = str(
            data["child5_name1"]
            + " "
            + data["child5_name2"]
            + " "
            + data["child5_name3"]
        )
        p3.add_run(
            f"                "
            f"- Свидетельство о рождении {fio}\n"
        )
    if data["lagguage_cost"] != "":
        p6.add_run(
            f"                "
            f"- Платежные документы на перевозку багажа\n"
        )
    if data["transfer_cost"] != "":
        p6.add_run(
            f"                "
            f"- Платежные документы на трансфер\n"
        )
    if data["visa_cost"] != "":
        p6.add_run(
            f"                "
            f"- Платежные документы на оформление миграционных документов (визы)\n"
        )

###############################    PARAGRAPH 7   ###############################
    fio = (
        data["rabotnik_name1"]
        + " "
        + data["rabotnik_name2"]
        + " "
        + data["rabotnik_name3"]
    )
    p7 = document.add_paragraph()
    p7.alignment = 2
    p7.paragraph_format.space_before = Pt(13)
    currdate = datetime.today().strftime("%d.%m.%Y")
    p7.add_run(f"                             / {fio} / {currdate}\n").underline = True
    tmp = p7.add_run(f"                 Подпись, расшифровака, дата          ")
    tmp.font.size = Pt(9)
    tmp.italic = True

###############################   PAGE BORDERS   ###############################
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(3.2)
        section.right_margin = Cm(1.7)

###############################      FOOTER      ###############################
'''
    section = document.sections[0]
    footer = section.footer
    table_3 = footer.add_table(0, 3, Cm(6))
    row_cells = table_3.add_row().cells
    row_cells[0].text = str("PL-01-04-06/05-01/01")
    row_cells[1].text = str(
        "Первый въезд иностранного работника "
        "(билеты, гостиница, виза, багаж, подъемные, "
        "переезд семьи, компенсация НДФЛ)"
    )
    row_cells[2].text = str(
        "Члены семьи работника: супруг/супруга, "
        "действующий законный брак, несовершеннолетние "
        "дети до 18 лет"
    )
    set_table_params_lo(
        table=table_3,
        widths=(Cm(3.5), Cm(7.5), Cm(5.8)),
        height=Cm(1.4),
        heightrule=WD_ROW_HEIGHT_RULE.EXACTLY,
        fontname="Nimbus Sans Narrow",
        fontsize=Pt(9),
        alignment=0,
        verticalalignment=WD_ALIGN_VERTICAL.CENTER,
    )
    table_3.allow_autofit = False
    table_3.autofit = False
    table_3.style = "Table Grid"
'''

###############################    SAVE FILE     ###############################
    try:
        os.makedirs(os.path.expanduser("./files/doc3/"))
    except:
        pass
    document.save(f"./files/doc3/document_pdf{user}.docx")

###############################  CONVERT TO PDF  ###############################
    subprocess.run(
        ["doc2pdf", f"./files/doc3/document_pdf{user}.docx"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
###############################    RENAME FILE   ###############################
    subprocess.run(
        ["mv", f"./files/doc3/document_pdf{user}.pdf", f"./files/doc3/document{user}.pdf"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        encoding="utf-8",
    )
###############################        END       ###############################
    return None
